import io
import pdfplumber
from docx import Document
from groq import AsyncGroq
from fastapi import HTTPException
from app.config import settings

client = AsyncGroq(api_key=settings.groq_api_key)

# Formatos suportados pelo Groq Whisper
AUDIO_FORMATOS_SUPORTADOS = {"mp3", "mp4", "mpeg", "mpga", "m4a", "wav", "ogg", "opus", "flac"}

# MIME types por extensão
AUDIO_MIMES = {
    "mp3": "audio/mpeg",
    "mp4": "audio/mp4",
    "wav": "audio/wav",
    "ogg": "audio/ogg",
    "opus": "audio/opus",
    "flac": "audio/flac",
    "m4a": "audio/m4a",
    "mpeg": "audio/mpeg",
    "mpga": "audio/mpeg",
    # webm não é suportado pelo Groq — tratado abaixo
    "webm": "audio/ogg",  # webm com codec opus → renomear para ogg funciona
}


def extrair_texto_pdf(conteudo: bytes) -> str:
    """Extrai texto de um PDF usando pdfplumber."""
    texto_total = []
    with pdfplumber.open(io.BytesIO(conteudo)) as pdf:
        for page in pdf.pages:
            texto = page.extract_text()
            if texto:
                texto_total.append(texto)
    return "\n\n".join(texto_total)


def extrair_texto_docx(conteudo: bytes) -> str:
    """Extrai texto de um arquivo DOCX."""
    doc = Document(io.BytesIO(conteudo))
    paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]

    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [cel.text.strip() for cel in linha.cells if cel.text.strip()]
            if celulas:
                paragrafos.append(" | ".join(celulas))

    return "\n".join(paragrafos)


async def transcrever_audio(conteudo: bytes, filename: str) -> str:
    """
    Transcreve áudio usando Groq Whisper.
    
    Suporta: mp3, mp4, wav, ogg, opus, flac, m4a
    Tratamento especial para webm (Chrome grava neste formato):
      → renomeia para .ogg pois o codec interno é opus, que o Groq aceita
    """
    if not filename:
        filename = "gravacao.ogg"

    # Determinar extensão
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "ogg"

    # webm do Chrome (codec opus) → tratado como ogg
    if ext == "webm":
        filename_envio = filename.rsplit(".", 1)[0] + ".ogg"
        mime = "audio/ogg"
    elif ext in AUDIO_MIMES:
        filename_envio = filename
        mime = AUDIO_MIMES[ext]
    else:
        raise HTTPException(
            status_code=422,
            detail=(
                f"Formato de áudio '{ext}' não é suportado. "
                "Use MP3, WAV, OGG, OPUS, M4A ou FLAC."
            ),
        )

    try:
        transcricao = await client.audio.transcriptions.create(
            file=(filename_envio, io.BytesIO(conteudo), mime),
            model="whisper-large-v3",
            language="pt",
            response_format="text",
        )

        # A resposta pode ser string direta ou objeto com .text
        texto = transcricao if isinstance(transcricao, str) else getattr(transcricao, "text", "")
        return texto.strip()

    except HTTPException:
        raise
    except Exception as e:
        err_str = str(e).lower()

        if "rate" in err_str or "429" in err_str or "limit" in err_str:
            raise HTTPException(
                status_code=429,
                detail="O serviço de transcrição atingiu o limite de requisições. Aguarde alguns minutos e tente novamente.",
            )
        if "format" in err_str or "unsupported" in err_str or "invalid" in err_str:
            raise HTTPException(
                status_code=422,
                detail=(
                    "O formato ou qualidade do áudio não foi reconhecido. "
                    "Tente gravar novamente em um ambiente silencioso, ou envie um arquivo MP3/WAV."
                ),
            )
        if "size" in err_str or "too large" in err_str:
            raise HTTPException(
                status_code=413,
                detail="O arquivo de áudio é muito grande. O limite é de 25 MB. Grave um trecho mais curto.",
            )

        raise HTTPException(
            status_code=500,
            detail=(
                "Não foi possível transcrever o áudio. "
                "Verifique a qualidade da gravação e tente novamente, "
                "ou use a opção de texto digitado."
            ),
        )
